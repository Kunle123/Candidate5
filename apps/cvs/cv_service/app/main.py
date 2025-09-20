from fastapi import FastAPI, HTTPException, Depends, status, Query, Body, Request, File, UploadFile, Path
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session
from sqlalchemy import desc, asc
import os
import uuid
import json
import jwt
from datetime import datetime
from typing import Optional, List, Dict, Any, Union
from fastapi.security import OAuth2PasswordBearer
import logging
import sys
from docx import Document
from tempfile import NamedTemporaryFile
from io import BytesIO
import time
import base64
import io
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
import httpx
from sqlalchemy.exc import NoResultFound
import traceback
import re

# Import database and models
from .database import get_db_session, is_sqlite, engine, Base
from . import models

# Configure logging
logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger("cv_service")

# Create FastAPI app
app = FastAPI(title="CandidateV CV Service")

# Environment variables
JWT_SECRET = os.getenv("JWT_SECRET", "development_secret_key")
JWT_ALGORITHM = os.getenv("JWT_ALGORITHM", "HS256")
CORS_ORIGINS = os.getenv("CORS_ORIGINS", "http://localhost:3000,http://localhost:5173,http://localhost:5174,http://localhost:5175,https://c5-frontend-pied.vercel.app").split(",")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"]
)

# OAuth2 scheme
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/api/auth/login", auto_error=False)

# Pydantic models
class CVMetadata(BaseModel):
    name: str
    description: Optional[str] = None
    is_default: bool = False
    version: int = 1
    last_modified: datetime = Field(default_factory=datetime.utcnow)

    class Config:
        orm_mode = True

class CVCreate(BaseModel):
    name: str
    description: Optional[str] = None
    is_default: bool = False
    template_id: str = "default"
    base_cv_id: Optional[str] = None  # If copying from existing CV

class CVUpdateMetadata(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    is_default: Optional[bool] = None

class CVUpdateContent(BaseModel):
    template_id: Optional[str] = None
    style_options: Optional[Dict[str, Any]] = None
    personal_info: Optional[Dict[str, Any]] = None
    summary: Optional[str] = None
    
    # We'll implement these relationships separately
    # experiences: Optional[List[Dict[str, Any]]] = None
    # education: Optional[List[Dict[str, Any]]] = None
    # skills: Optional[List[Dict[str, Any]]] = None
    # languages: Optional[List[Dict[str, Any]]] = None
    # projects: Optional[List[Dict[str, Any]]] = None
    # certifications: Optional[List[Dict[str, Any]]] = None
    # references: Optional[List[Dict[str, Any]]] = None
    
    custom_sections: Optional[Dict[str, Any]] = None

class ProfessionalCVFormatter:
    """Professional CV formatter with enhanced styling and layout capabilities."""
    def __init__(self):
        self.doc = Document()
        self.setup_document()
        self.setup_styles()
    def setup_document(self):
        section = self.doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    def setup_styles(self):
        self.colors = {
            'primary': RGBColor(44, 62, 80),
            'secondary': RGBColor(52, 152, 219),
            'accent': RGBColor(149, 165, 166),
            'text': RGBColor(44, 62, 80),
            'light_text': RGBColor(127, 140, 141)
        }
        self.font_sizes = {
            'name': 24,
            'title': 18,
            'section_heading': 14,
            'job_title': 12,
            'company': 11,
            'body': 10,
            'small': 9
        }
        self.spacing = {
            'section_before': Pt(20),
            'section_after': Pt(8),
            'item_before': Pt(6),
            'item_after': Pt(4),
            'line_spacing': 1.15
        }
    def set_font_style(self, paragraph, font_size, bold=False, color=None, italic=False):
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            run.font.size = Pt(font_size)
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = color
        paragraph.paragraph_format.line_spacing = self.spacing['line_spacing']
    def add_horizontal_line(self, color=None):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        p = para._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr, 'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr', 'w:sectPrChange')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        if color:
            bottom.set(qn('w:color'), f"{color[0]:02x}{color[1]:02x}{color[2]:02x}")
        else:
            accent_color = self.colors['accent']
            bottom.set(qn('w:color'), f"{accent_color[0]:02x}{accent_color[1]:02x}{accent_color[2]:02x}")
        pBdr.append(bottom)
    def add_header_section(self, name, title, contact_info):
        name_para = self.doc.add_paragraph(name)
        self.set_font_style(name_para, self.font_sizes['name'], bold=True, color=self.colors['primary'])
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        name_para.paragraph_format.space_after = Pt(4)
        title_para = self.doc.add_paragraph(title)
        self.set_font_style(title_para, self.font_sizes['title'], color=self.colors['secondary'])
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_para.paragraph_format.space_after = Pt(8)
        if contact_info:
            table = self.doc.add_table(rows=1, cols=len(contact_info))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, info in enumerate(contact_info):
                cell = table.cell(0, i)
                cell.text = info
                cell_para = cell.paragraphs[0]
                self.set_font_style(cell_para, self.font_sizes['small'], color=self.colors['light_text'])
                cell_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.add_horizontal_line()
    def add_section_heading(self, text):
        para = self.doc.add_paragraph()
        run = para.add_run(text.upper())
        self.set_font_style(para, self.font_sizes['section_heading'], bold=True, color=self.colors['primary'])
        para.paragraph_format.space_before = self.spacing['section_before']
        para.paragraph_format.space_after = self.spacing['section_after']
        return para
    def add_experience_block(self, title, company, location, dates, description=None):
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title)
        self.set_font_style(title_para, self.font_sizes['job_title'], bold=True, color=self.colors['text'])
        dates_run = title_para.add_run(f"\t{dates}")
        dates_run.font.name = 'Calibri'
        dates_run.font.size = Pt(self.font_sizes['small'])
        dates_run.font.color.rgb = self.colors['light_text']
        tab_stops = title_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_PARAGRAPH_ALIGNMENT.RIGHT)
        title_para.paragraph_format.space_before = self.spacing['item_before']
        title_para.paragraph_format.space_after = Pt(2)
        company_para = self.doc.add_paragraph(f"{company}")
        if location:
            company_para.add_run(f" • {location}")
        self.set_font_style(company_para, self.font_sizes['company'], italic=True, color=self.colors['secondary'])
        company_para.paragraph_format.space_after = Pt(4)
        if description:
            # Always treat as bullet points
            if isinstance(description, list):
                bullets = description
            else:
                # Split string by newlines, remove empty lines and manual bullets
                bullets = [line.lstrip('•').strip() for line in description.split('\n') if line.strip()]
            self.add_bullet_points(bullets)
    def add_education_block(self, degree, institution, location, year, details=None):
        degree_para = self.doc.add_paragraph()
        degree_run = degree_para.add_run(degree)
        self.set_font_style(degree_para, self.font_sizes['job_title'], bold=True, color=self.colors['text'])
        year_run = degree_para.add_run(f"\t{year}")
        year_run.font.name = 'Calibri'
        year_run.font.size = Pt(self.font_sizes['small'])
        year_run.font.color.rgb = self.colors['light_text']
        tab_stops = degree_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_PARAGRAPH_ALIGNMENT.RIGHT)
        degree_para.paragraph_format.space_before = self.spacing['item_before']
        degree_para.paragraph_format.space_after = Pt(2)
        inst_para = self.doc.add_paragraph(f"{institution}")
        if location:
            inst_para.add_run(f" • {location}")
        self.set_font_style(inst_para, self.font_sizes['company'], italic=True, color=self.colors['secondary'])
        inst_para.paragraph_format.space_after = Pt(4)
        if details:
            details_para = self.doc.add_paragraph(details)
            self.set_font_style(details_para, self.font_sizes['body'], color=self.colors['text'])
            details_para.paragraph_format.space_after = self.spacing['item_after']
            details_para.paragraph_format.left_indent = Inches(0.2)
    def add_skills_section(self, skills_dict):
        for category, skills in skills_dict.items():
            cat_para = self.doc.add_paragraph()
            cat_run = cat_para.add_run(f"{category}: ")
            self.set_font_style(cat_para, self.font_sizes['body'], bold=True, color=self.colors['text'])
            skills_run = cat_para.add_run(" • ".join(skills))
            skills_run.font.name = 'Calibri'
            skills_run.font.size = Pt(self.font_sizes['body'])
            skills_run.font.color.rgb = self.colors['text']
            cat_para.paragraph_format.space_after = Pt(4)
            cat_para.paragraph_format.left_indent = Inches(0.2)
    def add_bullet_points(self, items):
        for item in items:
            para = self.doc.add_paragraph(f"• {item}")
            self.set_font_style(para, self.font_sizes['body'], color=self.colors['text'])
            para.paragraph_format.left_indent = Inches(0.2)
            para.paragraph_format.space_after = Pt(3)
    def save_to_buffer(self):
        buf = BytesIO()
        self.doc.save(buf)
        return buf.getvalue()

# Helper function to verify JWT tokens
async def verify_token(token: Optional[str] = Depends(oauth2_scheme)):
    if token is None:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Not authenticated",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    try:
        # Verify token locally
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        
        # Accept both 'user_id' and 'id' fields
        user_id = payload.get("user_id") or payload.get("id")
        
        if user_id is None:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Invalid token format",
                headers={"WWW-Authenticate": "Bearer"},
            )
        
        return {"user_id": user_id}
    
    except (jwt.PyJWTError, Exception) as e:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail=f"Invalid authentication credentials: {str(e)}",
            headers={"WWW-Authenticate": "Bearer"},
        )

# Helper function to serialize database objects to JSON-compatible dictionaries
def serialize_cv(cv, include_relationships=True, db=None):
    """Convert a CV database object to a dictionary, with job_title, company, and cover letter info."""
    result = {
        "id": str(cv.id) if hasattr(cv.id, "hex") else cv.id,
        "user_id": str(cv.user_id) if hasattr(cv.user_id, "hex") else cv.user_id,
        "metadata": {
            "name": cv.name,
            "description": cv.description,
            "is_default": cv.is_default,
            "version": cv.version,
            "last_modified": cv.last_modified.isoformat() if cv.last_modified else None
        },
        "content": {
            "template_id": cv.template_id,
            "style_options": json.loads(cv.style_options) if isinstance(cv.style_options, str) else cv.style_options or {},
            "personal_info": json.loads(cv.personal_info) if isinstance(cv.personal_info, str) else cv.personal_info or {},
            "summary": cv.summary,
            "custom_sections": json.loads(cv.custom_sections) if isinstance(cv.custom_sections, str) else cv.custom_sections or {},
        },
        "created_at": cv.created_at.isoformat() if cv.created_at else None,
        "updated_at": cv.updated_at.isoformat() if cv.updated_at else None
    }
    # Extract job_title and company_name from personal_info if available
    personal_info = result["content"].get("personal_info", {})
    job_title = personal_info.get("job_title")
    company_name = personal_info.get("company_name") or personal_info.get("company")
    result["job_title"] = job_title
    result["company_name"] = company_name
    # Cover letter info (direct link)
    cover_letter_available = False
    cover_letter_download_url = None
    if getattr(cv, "cover_letter_id", None) and db is not None:
        from .models import CV as CVModel
        cover_letter = db.query(CVModel).filter(CVModel.id == cv.cover_letter_id, CVModel.type == "cover_letter").first()
        if cover_letter:
            cover_letter_available = True
            cover_letter_download_url = f"/api/cv/{cover_letter.id}/download"
    result["cover_letter_available"] = cover_letter_available
    result["cover_letter_download_url"] = cover_letter_download_url
    return result

class ApplicationHistoryIn(BaseModel):
    job_title: str
    company_name: str
    job_description: Optional[str] = None
    applied_at: Optional[str] = None
    salary: Optional[str] = None
    contact_name: Optional[str] = None
    contact_number: Optional[str] = None
    organisation: Optional[str] = None

class ApplicationHistoryOut(ApplicationHistoryIn):
    id: str
    created_at: Optional[str] = None

    class Config:
        orm_mode = True

def serialize_application_history(entry):
    return {
        "id": str(entry.id),
        "job_title": entry.job_title,
        "company_name": entry.company_name,
        "job_description": entry.job_description,
        "applied_at": entry.applied_at.isoformat() if entry.applied_at else None,
        "salary": entry.salary,
        "contact_name": entry.contact_name,
        "contact_number": entry.contact_number,
        "organisation": entry.organisation,
        "created_at": entry.created_at.isoformat() if entry.created_at else None,
    }

def generate_filename(name, filetype, company):
    initials = "".join([part[0] for part in name.split() if part]).upper() or "CV"
    company_clean = re.sub(r'[^A-Za-z0-9]', '', company or 'Company')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    suffix = "CV" if filetype == "cv" else "CLetter"
    return f"{initials}_{suffix}_{company_clean}_{timestamp}.docx"

# Routes
@app.get("/")
async def root():
    return {"message": "CandidateV CV Service"}

@app.get("/api/health")
async def health_check():
    """Health check endpoint for deployment verification"""
    import platform
    
    return {
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "version": "1.0.0",
        "database_type": "SQLite" if is_sqlite else "PostgreSQL",
        "environment": {
            "python_version": sys.version,
            "platform": platform.platform(),
            "cwd": os.getcwd()
        }
    }

@app.get("/api/cv")
async def get_cvs(
    auth: dict = Depends(verify_token), 
    db: Session = Depends(get_db_session),
    type: str = Query(None, description="Filter by type: 'cv' or 'cover_letter'")
):
    """Get all CVs for the current user, optionally filtered by type."""
    user_id = auth["user_id"]
    query = db.query(models.CV).filter(models.CV.user_id == user_id)
    if type:
        query = query.filter(models.CV.type == type)
    cvs = query.order_by(
        desc(models.CV.is_default),
        desc(models.CV.last_modified)
    ).all()
    result = [serialize_cv(cv, db=db) for cv in cvs]
    return result

@app.post("/api/cv/upload", status_code=status.HTTP_201_CREATED)
async def create_cv_from_file(
    file: UploadFile = File(...),
    auth: dict = Depends(verify_token),
    db: Session = Depends(get_db_session)
):
    """Create a new CV from an uploaded file (legacy endpoint)."""
    logger.info(f"Received file upload: {file.filename}, content type: {file.content_type}")
    user_id = auth["user_id"]
    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name
    try:
        doc = Document(tmp_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        extracted_text = "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error parsing .docx: {e}")
        raise HTTPException(status_code=400, detail="Failed to parse .docx file")
    finally:
        os.unlink(tmp_path)
    # ... rest of legacy logic ...
    # (unchanged)
    # ... existing code ...

@app.middleware("http")
async def log_slow_requests(request: Request, call_next):
    start_time = time.time()
    response = await call_next(request)
    duration = time.time() - start_time
    if duration > 2:  # seconds
        logger.warning(f"Slow request: {request.method} {request.url} took {duration:.2f}s")
    return response

@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    logger.error(
        f"Unhandled Exception: {exc}\n"
        f"Request: {request.method} {request.url}\n"
        f"Client: {request.client.host if request.client else 'unknown'}",
        exc_info=True
    )
    return JSONResponse(
        status_code=500,
        content={"detail": "Internal server error", "error": str(exc)},
    )

import httpx
USER_SERVICE_URL = os.getenv("USER_SERVICE_URL", "http://user-service:8000/api/user/profile")

async def inject_pii_placeholders(payload, auth_header):
    """
    Replace {{CANDIDATE_NAME}} and {{CONTACT_INFO}} placeholders in payload with actual PII from user profile.
    Accepts contact_info as either a string placeholder or an array containing the placeholder.
    """
    contact_info = payload.get("contact_info")
    is_contact_info_placeholder = (
        contact_info == "{{CONTACT_INFO}}" or
        (isinstance(contact_info, list) and len(contact_info) == 1 and contact_info[0] == "{{CONTACT_INFO}}")
    )
    if not (payload.get("name") == "{{CANDIDATE_NAME}}" or is_contact_info_placeholder):
        return payload  # No placeholders to replace
    async with httpx.AsyncClient() as client:
        resp = await client.get(
            USER_SERVICE_URL,
            headers={"Authorization": auth_header}
        )
        if resp.status_code != 200:
            raise HTTPException(status_code=500, detail="Failed to fetch user profile for PII injection.")
        user_profile = resp.json()
    # Replace name
    if payload.get("name") == "{{CANDIDATE_NAME}}":
        actual_name = user_profile.get("name")
        if not actual_name:
            raise HTTPException(status_code=400, detail="User profile missing name for PII injection.")
        payload["name"] = actual_name
        logger.info(f"[PII] Replaced CANDIDATE_NAME placeholder with: {payload['name']}")
    # Replace contact_info
    if is_contact_info_placeholder:
        address = user_profile.get("address_line1") or ""
        city = user_profile.get("city_state_postal") or ""
        email = user_profile.get("email") or ""
        phone = user_profile.get("phone_number") or ""
        linkedin = user_profile.get("linkedin") or ""
        if not (address or city or email or phone or linkedin):
            raise HTTPException(status_code=400, detail="User profile missing contact info for PII injection.")
        payload["contact_info"] = [
            address,
            city,
            f"{email} | {phone} | {linkedin}"
        ]
    return payload

def is_placeholder(val, placeholder):
    return isinstance(val, str) and val.strip() == placeholder

@app.post("/api/cv")
@app.post("/api/cv/")
async def persist_cv(
    payload: dict = Body(...),
    auth: dict = Depends(verify_token),
    request: Request = None,
    db: Session = Depends(get_db_session)
):
    # --- Logging ---
    logger.info(f"[CV PERSIST] Received payload: {json.dumps(payload)[:1000]}" if payload else "[CV PERSIST] Received empty payload!")
    logger.info(f"[CV PERSIST] User: {auth}")
    # --- Validation: allow placeholders ---
    required_fields = ["name", "experience"]
    missing = [
        f for f in required_fields
        if not payload.get(f)
        and not is_placeholder(payload.get(f), "{{CANDIDATE_NAME}}")
    ]
    if missing:
        logger.warning(f"[CV PERSIST] Missing required fields: {missing}")
        raise HTTPException(status_code=400, detail=f"Missing required fields: {missing}")
    if not isinstance(payload.get("experience"), list) or not payload["experience"]:
        logger.warning("[CV PERSIST] 'experience' must be a non-empty list")
        raise HTTPException(status_code=400, detail="'experience' must be a non-empty list")
    for idx, job in enumerate(payload["experience"]):
        if not isinstance(job, dict) or not job.get("job_title"):
            logger.warning(f"[CV PERSIST] Experience[{idx}] missing 'job_title'")
            raise HTTPException(status_code=400, detail=f"Experience[{idx}] missing 'job_title'")
    # --- PII Placeholder Replacement ---
    auth_header = request.headers.get("authorization") if request else None
    if not auth_header:
        raise HTTPException(status_code=401, detail="Missing Authorization header for PII injection.")
    payload = await inject_pii_placeholders(payload, auth_header)
    # --- Validation: after PII injection ---
    if not payload.get("name") or is_placeholder(payload.get("name"), "{{CANDIDATE_NAME}}"):
        logger.warning("[CV PERSIST] Name missing after PII injection.")
        raise HTTPException(status_code=400, detail="Name missing after PII injection.")
    # contact_info is now optional, but if present, must be an array of strings
    if "contact_info" in payload and payload["contact_info"] is not None:
        if not isinstance(payload["contact_info"], list) or not all(isinstance(x, str) for x in payload["contact_info"]):
            logger.warning("[CV PERSIST] contact_info must be an array of strings after PII injection.")
            raise HTTPException(status_code=400, detail="'contact_info' must be an array of strings after PII injection")
    # --- Existing logic ---
    try:
        logger.info(f"[DEBUG] Starting CV DOCX generation (hierarchical JSON)")
        cv = ProfessionalCVFormatter()
        # --- Use explicit structured fields ---
        name = payload.get("name", "")
        job_title = payload.get("job_title", "")
        contact_info = payload.get("contact_info", [])
        summary = extract_content(payload.get("summary", ""))
        core_competencies = extract_list_content(payload.get("core_competencies", []))
        experience = payload.get("experience", [])
        education = payload.get("education", [])
        certifications = extract_list_content(payload.get("certifications", []))
        include_keywords = payload.get("includeKeywords", False)
        include_relevant_experience = payload.get("includeRelevantExperience", False)
        logger.info(f"includeKeywords: {include_keywords}, includeRelevantExperience: {include_relevant_experience}")
        # Header
        cv.add_header_section(name, job_title, contact_info)
        # Summary
        if summary:
            cv.add_section_heading("Professional Summary")
            para = cv.doc.add_paragraph(summary)
            cv.set_font_style(para, cv.font_sizes['body'], color=cv.colors['text'])
            para.paragraph_format.space_after = Pt(12)
        # Relevant Achievements
        relevant_achievements = extract_list_content(payload.get("relevant_achievements", []))
        if relevant_achievements:
            cv.add_section_heading("Relevant Achievements")
            cv.add_bullet_points(relevant_achievements)
        # Core Competencies
        if core_competencies:
            cv.add_section_heading("Core Competencies")
            para = cv.doc.add_paragraph()
            para.add_run(" ".join([f"• {kw}" for kw in core_competencies]))
            cv.set_font_style(para, cv.font_sizes['body'], color=cv.colors['text'])
            para.paragraph_format.space_after = Pt(8)
        # Experience
        if experience:
            cv.add_section_heading("Professional Experience")
            for job in experience:
                responsibilities = job.get("responsibilities")
                if responsibilities is None:
                    responsibilities = job.get("bullets", [])
                responsibilities = extract_list_content(responsibilities)
                cv.add_experience_block(
                    title=job.get("job_title", "") or job.get("title", ""),
                    company=job.get("company_name", "") or job.get("company", ""),
                    location=job.get("location", ""),
                    dates=job.get("dates", "") or f"{job.get('start_date', '')} – {job.get('end_date', '')}",
                    description=responsibilities
                )
        # Education
        if education:
            cv.add_section_heading("Education")
            for edu in education:
                cv.add_education_block(
                    degree=edu.get("degree", ""),
                    institution=edu.get("institution", ""),
                    location=edu.get("location", ""),
                    year=edu.get("dates", "")
                )
        # Certifications
        if certifications:
            cv.add_section_heading("Certifications")
            cv.add_bullet_points(certifications)
        docx_bytes = cv.save_to_buffer()
        logger.info(f"[DEBUG] DOCX bytes length: {len(docx_bytes)}")
        company_name = payload.get("company_name")
        personal_info = {}
        if job_title:
            personal_info["job_title"] = job_title
        if company_name:
            personal_info["company"] = company_name
        cover_letter_id = None
        if payload.get("cover_letter"):
            cover_letter_text = extract_content(payload.get("cover_letter", ""))
            cover_doc = Document()
            cover_doc.add_heading("Cover Letter", 0)
            for line in cover_letter_text.splitlines():
                if line.strip() == "":
                    cover_doc.add_paragraph()
                else:
                    para = cover_doc.add_paragraph(line)
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                        run.font.size = Pt(11)
            cover_buf = BytesIO()
            cover_doc.save(cover_buf)
            cover_buf.seek(0)
            cover_docx_bytes = cover_buf.getvalue()
            cover_personal_info = {}
            if job_title:
                cover_personal_info["job_title"] = job_title
            if company_name:
                cover_personal_info["company"] = company_name
            from .models import CV
            cover_letter_obj = CV(
                id=uuid.uuid4(),
                user_id=auth["user_id"],
                name="Generated Cover Letter",
                description="Cover letter generated via API",
                is_default=False,
                version=1,
                template_id="default",
                summary=None,
                docx_file=cover_docx_bytes,
                type="cover_letter",
                cover_letter_id=None,
                personal_info=json.dumps(cover_personal_info) if cover_personal_info else None
            )
            db.add(cover_letter_obj)
            db.commit()
            db.refresh(cover_letter_obj)
            cover_letter_id = str(cover_letter_obj.id)
            logger.info(f"[DEBUG] Cover letter persisted with id: {cover_letter_id}")
        from .models import CV
        logger.info("[DEBUG] Persisting CV to DB")
        new_cv = CV(
            id=uuid.uuid4(),
            user_id=auth["user_id"],
            name="Generated CV",
            description="CV generated via API",
            is_default=False,
            version=1,
            template_id="default",
            summary=None,
            docx_file=docx_bytes,
            type="cv",
            cover_letter_id=cover_letter_id,
            personal_info=json.dumps(personal_info) if personal_info else None
        )
        db.add(new_cv)
        db.commit()
        db.refresh(new_cv)
        logger.info(f"[DEBUG] CV DOCX generated and persisted in DB for user_id={auth.get('user_id')}, cv_id={new_cv.id}")
        try:
            docx_b64 = base64.b64encode(docx_bytes).decode('utf-8')
        except Exception as e:
            logger.error(f"Base64 encoding failed: {e}", exc_info=True)
            raise HTTPException(status_code=500, detail="Failed to encode DOCX as base64")
        # Example priority assignment logic (replace with real logic as needed)
        def assign_priority(items, max_priority=4):
            # Assign priorities: 1 for first 2, 2 for next 2, etc.
            out = []
            for i, item in enumerate(items):
                prio = 1 + (i // 2)
                out.append({"content": item, "priority": min(prio, max_priority)})
            return out

        # Wrap summary and cover_letter
        summary_obj = {"content": summary, "priority": 1}
        cover_letter_obj = {"content": payload.get("cover_letter", ""), "priority": 1}

        # Wrap relevant_achievements
        relevant_achievements_obj = []
        if include_relevant_experience and payload.get("relevant_achievements"):
            relevant_achievements_obj = assign_priority(payload["relevant_achievements"])

        # Wrap core_competencies
        core_competencies_obj = []
        if include_keywords and core_competencies:
            core_competencies_obj = assign_priority(core_competencies)

        # Wrap certifications
        certifications_obj = assign_priority(certifications)

        # Wrap experience responsibilities
        experience_out = []
        for job in experience:
            responsibilities = job.get("responsibilities")
            if responsibilities is None:
                responsibilities = job.get("bullets", [])
            responsibilities_obj = assign_priority(responsibilities)
            experience_out.append({
                "job_title": job.get("job_title", "") or job.get("title", ""),
                "company_name": job.get("company_name", "") or job.get("company", ""),
                "dates": job.get("dates", "") or f"{job.get('start_date', '')} – {job.get('end_date', '')}",
                "responsibilities": responsibilities_obj
            })

        # Wrap education
        education_out = []
        for i, edu in enumerate(education):
            education_out.append({
                "degree": edu.get("degree", ""),
                "institution": edu.get("institution", ""),
                "year": edu.get("year", ""),
                "priority": 1 if i == 0 else 2
            })

        # Trimming guide
        trimming_guide = {
            "2_page_version": "Keep priority 1-2 content only",
            "3_page_version": "Keep priority 1-3 content only",
            "4_page_version": "Keep all content (priority 1-5)"
        }

        # Build response dict
        response = {
            "name": "{{CANDIDATE_NAME}}",
            "contact_info": ["{{CONTACT_INFO}}"],
            "summary": summary_obj,
            "experience": experience_out,
            "education": education_out,
            "certifications": certifications_obj,
            "cover_letter": cover_letter_obj,
            "job_title": job_title,
            "company_name": payload.get("company_name", ""),
            "trimming_guide": trimming_guide,
            "relevant_achievements": relevant_achievements_obj  # Always include, even if empty
        }
        if core_competencies_obj:
            response["core_competencies"] = core_competencies_obj
        if payload.get("salary"):
            response["salary"] = payload["salary"]
        if payload.get("contact_name"):
            response["contact_name"] = payload["contact_name"]
        if payload.get("contact_number"):
            response["contact_number"] = payload["contact_number"]
        response.pop("thread_id", None)
        response.pop("skills", None)
        response.pop("achievements", None)
        return JSONResponse(content=response)
    except Exception as e:
        logger.error(f"Error generating CV DOCX: {e}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail="Failed to generate CV DOCX")

@app.post("/api/cover-letter")
@app.post("/api/cover-letter/")
async def generate_cover_letter_docx(
    payload: dict = Body(...),
    auth: dict = Depends(verify_token),
    request: Request = None,
    db: Session = Depends(get_db_session)
):
    """
    Generate a Cover Letter DOCX from provided text and return as base64-encoded string in JSON.
    Stores job_title and company_name in personal_info if provided.
    Returns the new cover letter's ID for linking.
    """
    import base64
    try:
        logger.info(f"Received {request.method} to {request.url} from {request.client.host if request.client else 'unknown'} (Cover Letter only)")
        cover_letter = payload.get("cover_letter")
        if not cover_letter:
            logger.warning("Missing 'cover_letter' field in request body")
            raise HTTPException(status_code=400, detail="'cover_letter' field is required in the request body.")
        doc = Document()
        doc.add_heading("Cover Letter", 0)
        for line in cover_letter.splitlines():
            if line.strip() == "":
                doc.add_paragraph()
            else:
                doc.add_paragraph(line)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        docx_bytes = buf.getvalue()
        # Parse job_title and company_name from payload (from assistant)
        job_title = payload.get("job_title")
        company_name = payload.get("company_name")
        personal_info = {}
        if job_title:
            personal_info["job_title"] = job_title
        if company_name:
            personal_info["company"] = company_name
        # Persist to DB (as a separate CV record for now)
        from .models import CV
        new_cv = CV(
            id=uuid.uuid4(),
            user_id=auth["user_id"],
            name="Generated Cover Letter",
            description="Cover letter generated via API",
            is_default=False,
            version=1,
            template_id="default",
            summary=None,
            docx_file=docx_bytes,
            type="cover_letter",
            personal_info=json.dumps(personal_info) if personal_info else None
        )
        db.add(new_cv)
        db.commit()
        db.refresh(new_cv)
        logger.info(f"Cover Letter DOCX generated and persisted in DB for user_id={auth.get('user_id')}, cv_id={new_cv.id}")
        try:
            docx_b64 = base64.b64encode(docx_bytes).decode('utf-8')
        except Exception as e:
            logger.error(f"Base64 encoding failed: {e}", exc_info=True)
            raise HTTPException(status_code=500, detail="Failed to encode DOCX as base64")
        return {
            "filename": f"cover_letter_{new_cv.id}.docx",
            "filedata": docx_b64,
            "cv_id": str(new_cv.id)  # This is the cover letter's ID
        }
    except Exception as e:
        logger.error(f"Error generating Cover Letter DOCX: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to generate Cover Letter DOCX")

@app.get("/api/cv/{cv_id}")
async def get_cv(
    cv_id: str,
    auth: dict = Depends(verify_token),
    db: Session = Depends(get_db_session)
):
    """Get a specific CV."""
    user_id = auth["user_id"]
    logger.info(f"Attempting to get CV. ID: '{cv_id}', User ID: '{user_id}'")
    cv = None
    try:
        # Query CV
        cv = db.query(models.CV).filter(
            models.CV.id == cv_id,
            models.CV.user_id == user_id
        ).first()
        logger.info(f"Database query executed for CV ID: '{cv_id}'")
    except Exception as e:
        logger.error(f"Database query failed for CV ID: '{cv_id}'. Error: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Database error while fetching CV"
        )

    if not cv:
        logger.warning(f"CV not found or access denied for CV ID: '{cv_id}', User ID: '{user_id}'")
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="CV not found"
        )
    
    logger.info(f"CV found for ID: '{cv_id}'. Serializing.")
    try:
        serialized_cv = serialize_cv(cv, db=db)
        logger.info(f"Serialization successful for CV ID: '{cv_id}'")
        return serialized_cv
    except Exception as e:
        logger.error(f"Serialization failed for CV ID: '{cv_id}'. Error: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Error processing CV data"
        )

@app.put("/api/cv/{cv_id}/metadata")
async def update_cv_metadata(
    cv_id: str,
    metadata: CVUpdateMetadata,
    auth: dict = Depends(verify_token),
    db: Session = Depends(get_db_session)
):
    """Update CV metadata."""
    user_id = auth["user_id"]
    
    # Query CV
    cv = db.query(models.CV).filter(
        models.CV.id == cv_id,
        models.CV.user_id == user_id
    ).first()
    
    if not cv:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="CV not found"
        )
    
    # Update fields if provided
    if metadata.name is not None:
        cv.name = metadata.name
    
    if metadata.description is not None:
        cv.description = metadata.description
    
    if metadata.is_default is not None:
        # If this is the default CV, update other CVs
        if metadata.is_default:
            # Find all default CVs for this user and set them to non-default
            default_cvs = db.query(models.CV).filter(
                models.CV.user_id == user_id,
                models.CV.is_default == True,
                models.CV.id != cv_id
            ).all()
            
            for other_cv in default_cvs:
                other_cv.is_default = False
        
        cv.is_default = metadata.is_default
    
    # Update version and timestamps
    cv.version += 1
    cv.last_modified = datetime.utcnow()
    cv.updated_at = datetime.utcnow()
    
    # Save changes
    db.commit()
    db.refresh(cv)
    
    return serialize_cv(cv)

@app.put("/api/cv/{cv_id}/content")
async def update_cv_content(
    cv_id: str,
    content: CVUpdateContent,
    auth: dict = Depends(verify_token),
    db: Session = Depends(get_db_session)
):
    """Update CV content."""
    user_id = auth["user_id"]
    
    # Query CV
    cv = db.query(models.CV).filter(
        models.CV.id == cv_id,
        models.CV.user_id == user_id
    ).first()
    
    if not cv:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="CV not found"
        )
    
    # Update fields if provided
    if content.template_id is not None:
        cv.template_id = content.template_id
    
    if content.style_options is not None:
        cv.style_options = json.dumps(content.style_options) if is_sqlite else content.style_options
    
    if content.personal_info is not None:
        cv.personal_info = json.dumps(content.personal_info) if is_sqlite else content.personal_info
    
    if content.summary is not None:
        cv.summary = content.summary
    
    if content.custom_sections is not None:
        cv.custom_sections = json.dumps(content.custom_sections) if is_sqlite else content.custom_sections
    
    # Update version and timestamps
    cv.version += 1
    cv.last_modified = datetime.utcnow()
    cv.updated_at = datetime.utcnow()
    
    # Save changes
    db.commit()
    db.refresh(cv)
    
    return serialize_cv(cv)

@app.delete("/api/cv/{cv_id}")
async def delete_cv(
    cv_id: str,
    auth: dict = Depends(verify_token),
    db: Session = Depends(get_db_session)
):
    """Delete a CV."""
    user_id = auth["user_id"]
    
    # Query CV
    cv = db.query(models.CV).filter(
        models.CV.id == cv_id,
        models.CV.user_id == user_id
    ).first()
    
    if not cv:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="CV not found"
        )
    
    # Delete CV
    db.delete(cv)
    db.commit()
    
    return {"message": "CV deleted successfully"}

# --- Database Table Creation --- 
# Ensure tables are created on startup
# NOTE: In production, using Alembic for migrations is recommended
@app.on_event("startup")
def startup_event():
    logger.info("Running startup event: Creating database tables if they don't exist...")
    try:
        Base.metadata.create_all(bind=engine)
        logger.info("Database tables checked/created successfully.")
    except Exception as e:
        logger.error(f"Error creating database tables: {e}", exc_info=True)
        # Depending on the error, you might want to prevent startup
# --- End Database Table Creation --- 

# --- /cvs endpoints (spec-compliant) ---
@app.get("/cvs")
async def get_cvs_v2(auth: dict = Depends(verify_token), db: Session = Depends(get_db_session)):
    return await get_cvs(auth, db)

@app.post("/cvs", status_code=status.HTTP_201_CREATED)
async def create_cv_v2(file: UploadFile = File(...), auth: dict = Depends(verify_token), db: Session = Depends(get_db_session)):
    return await create_cv_from_file(file, auth, db) # Reusing the existing function

@app.get("/cvs/{cv_id}")
async def get_cv_v2(cv_id: str, auth: dict = Depends(verify_token), db: Session = Depends(get_db_session)):
    return await get_cv(cv_id, auth, db)

@app.put("/cvs/{cv_id}")
async def update_cv_content_v2(cv_id: str, content: CVUpdateContent, auth: dict = Depends(verify_token), db: Session = Depends(get_db_session)):
    return await update_cv_content(cv_id, content, auth, db)

@app.delete("/cvs/{cv_id}")
async def delete_cv_v2(cv_id: str, auth: dict = Depends(verify_token), db: Session = Depends(get_db_session)):
    return await delete_cv(cv_id, auth, db)

@app.post("/cvs/{cv_id}/analyze")
async def analyze_cv(cv_id: str, jobDescription: dict = Body(...), auth: dict = Depends(verify_token), db: Session = Depends(get_db_session)):
    # Placeholder: return mock missing keywords
    # In a real implementation, analyze the CV content and compare to jobDescription
    return {"missingKeywords": ["Python", "Leadership", "Teamwork"]}

@app.get("/cvs/{cv_id}/download")
async def download_cv(cv_id: str, auth: dict = Depends(verify_token), db: Session = Depends(get_db_session)):
    from .models import CV
    user_id = auth["user_id"]
    cv = db.query(CV).filter(CV.id == cv_id, CV.user_id == user_id).first()
    if not cv or not cv.docx_file:
        raise HTTPException(status_code=404, detail="CV or DOCX file not found")
    return StreamingResponse(
        io.BytesIO(cv.docx_file),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="cv_{cv_id}.docx"'
        }
    )

@app.get("/api/cv/{cv_id}/download")
async def download_persisted_docx(cv_id: str, auth: dict = Depends(verify_token), db: Session = Depends(get_db_session)):
    """
    Return the persisted DOCX as a base64-encoded string in JSON (proxy-friendly).
    """
    from .models import CV
    user_id = auth["user_id"]
    try:
        cv = db.query(CV).filter(CV.id == cv_id, CV.user_id == user_id).first()
        if not cv or not cv.docx_file:
            logger.error(f"CV or DOCX file not found for cv_id={cv_id}, user_id={user_id}")
            raise HTTPException(status_code=404, detail="CV or DOCX file not found")
        # Generate filename
        name = cv.name or "CV"
        filetype = cv.type or "cv"
        # Try to get company from personal_info
        company = None
        try:
            if cv.personal_info:
                info = json.loads(cv.personal_info) if isinstance(cv.personal_info, str) else cv.personal_info
                company = info.get("company") or info.get("company_name")
        except Exception as e:
            logger.warning(f"Error parsing personal_info for cv_id={cv_id}: {e}")
        filename = generate_filename(name, filetype, company)
        return StreamingResponse(
            io.BytesIO(cv.docx_file),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )
    except Exception as e:
        logger.error(f"Error in /api/cv/{{cv_id}}/download: {e}", exc_info=True)
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail="Failed to download CV DOCX")

@app.post("/api/cv/generate-docx")
async def generate_docx_from_json(
    payload: dict = Body(...),
    auth: dict = Depends(verify_token),
    request: Request = None
):
    # --- Logging ---
    logger.info(f"[DOCX GEN] Received payload: {json.dumps(payload)[:1000]}" if payload else "[DOCX GEN] Received empty payload!")
    # --- Validation: allow placeholders ---
    required_fields = ["name", "experience"]
    missing = [
        f for f in required_fields
        if not payload.get(f)
        and not is_placeholder(payload.get(f), "{{CANDIDATE_NAME}}")
    ]
    if missing:
        logger.warning(f"[DOCX GEN] Missing required fields: {missing}")
        raise HTTPException(status_code=400, detail=f"Missing required fields: {missing}")
    if not isinstance(payload.get("experience"), list) or not payload["experience"]:
        logger.warning("[DOCX GEN] 'experience' must be a non-empty list")
        raise HTTPException(status_code=400, detail="'experience' must be a non-empty list")
    for idx, job in enumerate(payload["experience"]):
        if not isinstance(job, dict) or not job.get("job_title"):
            logger.warning(f"[DOCX GEN] Experience[{idx}] missing 'job_title'")
            raise HTTPException(status_code=400, detail=f"Experience[{idx}] missing 'job_title'")
    # --- PII Placeholder Replacement ---
    auth_header = request.headers.get("authorization") if request else None
    if not auth_header:
        raise HTTPException(status_code=401, detail="Missing Authorization header for PII injection.")
    payload = await inject_pii_placeholders(payload, auth_header)
    # --- Validation: after PII injection ---
    if not payload.get("name") or is_placeholder(payload.get("name"), "{{CANDIDATE_NAME}}"):
        logger.warning("[DOCX GEN] Name missing after PII injection.")
        raise HTTPException(status_code=400, detail="Name missing after PII injection.")
    # contact_info is now optional, but if present, must be an array of strings
    if "contact_info" in payload and payload["contact_info"] is not None:
        if not isinstance(payload["contact_info"], list) or not all(isinstance(x, str) for x in payload["contact_info"]):
            logger.warning("[DOCX GEN] contact_info must be an array of strings after PII injection.")
            raise HTTPException(status_code=400, detail="'contact_info' must be an array of strings after PII injection")
    try:
        cv = ProfessionalCVFormatter()
        name = payload.get("name", "")
        job_title = payload.get("job_title", "")
        contact_info = payload.get("contact_info", [])
        summary = extract_content(payload.get("summary", ""))
        cover_letter = extract_content(payload.get("cover_letter", ""))
        core_competencies = extract_list_content(payload.get("core_competencies", []))
        relevant_achievements = extract_list_content(payload.get("relevant_achievements", []))
        experience = payload.get("experience", [])
        education = payload.get("education", [])
        certifications = extract_list_content(payload.get("certifications", []))
        # Header
        cv.add_header_section(name, job_title, contact_info)
        # Summary
        if summary:
            cv.add_section_heading("Professional Summary")
            para = cv.doc.add_paragraph(summary)
            cv.set_font_style(para, cv.font_sizes['body'], color=cv.colors['text'])
            para.paragraph_format.space_after = Pt(12)
        # Relevant Achievements
        relevant_achievements = extract_list_content(payload.get("relevant_achievements", []))
        if relevant_achievements:
            cv.add_section_heading("Relevant Achievements")
            cv.add_bullet_points(relevant_achievements)
        # Core Competencies
        if core_competencies:
            cv.add_section_heading("Core Competencies")
            para = cv.doc.add_paragraph()
            para.add_run(" ".join([f"• {kw}" for kw in core_competencies]))
            cv.set_font_style(para, cv.font_sizes['body'], color=cv.colors['text'])
            para.paragraph_format.space_after = Pt(8)
        # Experience
        if experience:
            cv.add_section_heading("Professional Experience")
            for job in experience:
                responsibilities = job.get("responsibilities")
                if responsibilities is None:
                    responsibilities = job.get("bullets", [])
                responsibilities = extract_list_content(responsibilities)
                cv.add_experience_block(
                    title=job.get("job_title", "") or job.get("title", ""),
                    company=job.get("company_name", "") or job.get("company", ""),
                    location=job.get("location", ""),
                    dates=job.get("dates", "") or f"{job.get('start_date', '')} – {job.get('end_date', '')}",
                    description=responsibilities
                )
        # Education
        if education:
            cv.add_section_heading("Education")
            for edu in education:
                cv.add_education_block(
                    degree=edu.get("degree", ""),
                    institution=edu.get("institution", ""),
                    location=edu.get("location", ""),
                    year=edu.get("dates", "")
                )
        # Certifications
        if certifications:
            cv.add_section_heading("Certifications")
            cv.add_bullet_points(certifications)
        docx_bytes = cv.save_to_buffer()
        import base64
        docx_b64 = base64.b64encode(docx_bytes).decode('utf-8')
        # Optionally handle cover letter
        cover_letter_b64 = None
        if cover_letter:
            cover_doc = Document()
            cover_doc.add_heading("Cover Letter", 0)
            for line in cover_letter.splitlines():
                if line.strip() == "":
                    cover_doc.add_paragraph()
                else:
                    para = cover_doc.add_paragraph(line)
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                        run.font.size = Pt(11)
            from io import BytesIO
            cover_buf = BytesIO()
            cover_doc.save(cover_buf)
            cover_buf.seek(0)
            cover_docx_bytes = cover_buf.getvalue()
            cover_letter_b64 = base64.b64encode(cover_docx_bytes).decode('utf-8')
        # Build response
        response = {
            "cv": docx_b64,
            "cover_letter": cover_letter_b64
        }
        # Optional fields
        for opt in ["salary", "contact_name", "contact_number", "company_name", "job_title"]:
            if payload.get(opt):
                response[opt] = payload[opt]
        return response
    except Exception as e:
        logger.error(f"Error generating DOCX from JSON: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail="Failed to generate DOCX from JSON")

@app.post("/api/applications")
async def create_application(
    payload: dict = Body(...),
    auth: dict = Depends(verify_token),
    db: Session = Depends(get_db_session),
    request: Request = None
):
    """
    Create a single ApplicationHistory record for a job application.
    """
    from .models import ApplicationHistory
    user_id = auth["user_id"]
    job_title = payload.get("role_title") or payload.get("job_title")
    company_name = payload.get("company_name") or payload.get("company") or payload.get("organisation")
    contact_name = payload.get("contact_name")
    contact_number = payload.get("contact_number")
    salary = payload.get("salary")
    applied_at = payload.get("applied_at")
    job_description = payload.get("job_description")
    entry = ApplicationHistory(
        user_id=user_id,
        job_title=job_title,
        company_name=company_name,
        contact_name=contact_name,
        contact_number=contact_number,
        salary=salary,
        applied_at=applied_at,
        job_description=job_description
    )
    db.add(entry)
    db.commit()
    db.refresh(entry)
    return serialize_application_history(entry)

@app.get("/api/applications")
async def list_applications(auth: dict = Depends(verify_token), db: Session = Depends(get_db_session)):
    from .models import Application
    user_id = auth["user_id"]
    apps = db.query(Application).filter(Application.user_id == user_id).order_by(Application.created_at.desc()).all()
    return [
        {
            "id": str(app.id),
            "role_title": app.role_title,
            "created_at": app.created_at.isoformat(),
        } for app in apps
    ]

@app.get("/api/applications/{id}/cv")
async def download_application_cv(id: str = Path(...), auth: dict = Depends(verify_token), db: Session = Depends(get_db_session)):
    from .models import Application
    user_id = auth["user_id"]
    app = db.query(Application).filter(Application.id == id, Application.user_id == user_id).first()
    if not app or not app.cv_docx_file:
        raise HTTPException(status_code=404, detail="Application CV or DOCX file not found")
    # Generate filename (fallbacks if info missing)
    name = getattr(app, "candidate_name", "CV")
    company = getattr(app, "company_name", None)
    filename = generate_filename(name, "cv", company)
    return StreamingResponse(
        io.BytesIO(app.cv_docx_file),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        }
    )

@app.get("/api/applications/{id}/cover-letter")
async def download_application_cover_letter(id: str = Path(...), auth: dict = Depends(verify_token), db: Session = Depends(get_db_session)):
    from .models import Application
    user_id = auth["user_id"]
    app = db.query(Application).filter(Application.id == id, Application.user_id == user_id).first()
    if not app or not app.cover_letter_docx_file:
        raise HTTPException(status_code=404, detail="Application cover letter or DOCX file not found")
    return StreamingResponse(
        io.BytesIO(app.cover_letter_docx_file),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="application_cover_letter_{id}.docx"'
        }
    )

@app.get("/api/application-history", response_model=List[ApplicationHistoryOut])
async def get_application_history(auth: dict = Depends(verify_token), db: Session = Depends(get_db_session)):
    user_id = auth["user_id"]
    entries = db.query(models.ApplicationHistory).filter_by(user_id=user_id).order_by(models.ApplicationHistory.created_at.desc()).all()
    return [serialize_application_history(e) for e in entries]

@app.post("/api/application-history", response_model=ApplicationHistoryOut)
async def create_application_history(payload: ApplicationHistoryIn, auth: dict = Depends(verify_token), db: Session = Depends(get_db_session)):
    user_id = auth["user_id"]
    entry = models.ApplicationHistory(user_id=user_id, **payload.dict())
    db.add(entry)
    db.commit()
    db.refresh(entry)
    return serialize_application_history(entry)

@app.patch("/api/application-history/{id}", response_model=ApplicationHistoryOut)
async def update_application_history(id: str, payload: ApplicationHistoryIn, auth: dict = Depends(verify_token), db: Session = Depends(get_db_session)):
    user_id = auth["user_id"]
    entry = db.query(models.ApplicationHistory).filter_by(id=id, user_id=user_id).first()
    if not entry:
        raise HTTPException(status_code=404, detail="Not found")
    for k, v in payload.dict(exclude_unset=True).items():
        setattr(entry, k, v)
    db.commit()
    db.refresh(entry)
    return serialize_application_history(entry)

@app.delete("/api/application-history/{id}")
async def delete_application_history(id: str, auth: dict = Depends(verify_token), db: Session = Depends(get_db_session)):
    user_id = auth["user_id"]
    entry = db.query(models.ApplicationHistory).filter_by(id=id, user_id=user_id).first()
    if not entry:
        raise HTTPException(status_code=404, detail="Not found")
    db.delete(entry)
    db.commit()
    return {"success": True}
