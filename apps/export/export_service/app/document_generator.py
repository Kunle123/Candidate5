import os
import asyncio
import logging
import jinja2
import tempfile
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, Tuple

# Import PDF and DOCX generation libraries
from weasyprint import HTML, CSS
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Set up logging
logger = logging.getLogger("export_service.document_generator")

# Templates directory
TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")
os.makedirs(TEMPLATES_DIR, exist_ok=True)

# Default template paths
DEFAULT_TEMPLATE_PATH = os.path.join(TEMPLATES_DIR, "default.html")

# Create default template if it doesn't exist
if not os.path.exists(DEFAULT_TEMPLATE_PATH):
    with open(DEFAULT_TEMPLATE_PATH, "w") as f:
        f.write("""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ cv.title }}</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            margin: 0;
            padding: 20px;
            color: #333;
        }
        .header {
            text-align: center;
            margin-bottom: 20px;
        }
        h1 {
            color: #2c3e50;
            margin-bottom: 5px;
        }
        .personal-info {
            text-align: center;
            margin-bottom: 20px;
        }
        .section {
            margin-bottom: 20px;
        }
        .section-title {
            border-bottom: 2px solid #3498db;
            padding-bottom: 5px;
            margin-bottom: 10px;
            color: #2c3e50;
        }
        .experience-item, .education-item {
            margin-bottom: 15px;
        }
        .item-header {
            display: flex;
            justify-content: space-between;
            margin-bottom: 5px;
        }
        .item-title {
            font-weight: bold;
        }
        .item-date {
            color: #7f8c8d;
        }
        .item-subtitle {
            font-style: italic;
            margin-bottom: 5px;
        }
        .item-description {
            margin-top: 5px;
        }
        .skills-list {
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
        }
        .skill-item {
            background-color: #e6f3ff;
            padding: 5px 10px;
            border-radius: 15px;
        }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ cv.title }}</h1>
    </div>
    
    <div class="personal-info">
        <div>{{ cv.personal_info.first_name }} {{ cv.personal_info.last_name }}</div>
        <div>{{ cv.personal_info.email }} | {{ cv.personal_info.phone }}</div>
        <div>{{ cv.personal_info.location }}</div>
    </div>
    
    {% if cv.summary %}
    <div class="section">
        <h2 class="section-title">Summary</h2>
        <p>{{ cv.summary }}</p>
    </div>
    {% endif %}
    
    {% if cv.experience and cv.experience|length > 0 %}
    <div class="section">
        <h2 class="section-title">Experience</h2>
        {% for exp in cv.experience %}
        <div class="experience-item">
            <div class="item-header">
                <div class="item-title">{{ exp.title }}</div>
                <div class="item-date">
                    {{ exp.start_date }} - {% if exp.current %}Present{% else %}{{ exp.end_date }}{% endif %}
                </div>
            </div>
            <div class="item-subtitle">{{ exp.company }}{% if exp.location %}, {{ exp.location }}{% endif %}</div>
            <div class="item-description">{{ exp.description }}</div>
        </div>
        {% endfor %}
    </div>
    {% endif %}
    
    {% if cv.education and cv.education|length > 0 %}
    <div class="section">
        <h2 class="section-title">Education</h2>
        {% for edu in cv.education %}
        <div class="education-item">
            <div class="item-header">
                <div class="item-title">{{ edu.degree }} in {{ edu.field_of_study }}</div>
                <div class="item-date">
                    {{ edu.start_date }} - {{ edu.end_date }}
                </div>
            </div>
            <div class="item-subtitle">{{ edu.institution }}</div>
            {% if edu.description %}
            <div class="item-description">{{ edu.description }}</div>
            {% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}
    
    {% if cv.skills and cv.skills|length > 0 %}
    <div class="section">
        <h2 class="section-title">Skills</h2>
        <div class="skills-list">
            {% for skill in cv.skills %}
            <div class="skill-item">{{ skill }}</div>
            {% endfor %}
        </div>
    </div>
    {% endif %}
    
    {% if cv.certifications and cv.certifications|length > 0 %}
    <div class="section">
        <h2 class="section-title">Certifications</h2>
        {% for cert in cv.certifications %}
        <div class="education-item">
            <div class="item-header">
                <div class="item-title">{{ cert.name }}</div>
                <div class="item-date">{{ cert.date }}</div>
            </div>
            <div class="item-subtitle">{{ cert.issuer }}</div>
            {% if cert.description %}
            <div class="item-description">{{ cert.description }}</div>
            {% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}
</body>
</html>
""")

# Initialize Jinja2 environment
template_loader = jinja2.FileSystemLoader(searchpath=TEMPLATES_DIR)
template_env = jinja2.Environment(loader=template_loader)

class DocumentGenerator:
    """
    Class for generating PDF and DOCX documents from CV data
    """
    
    @staticmethod
    async def generate_pdf_weasyprint(cv_data: Dict[str, Any], output_path: str, template_options: Optional[Dict[str, Any]] = None) -> str:
        """Generate a PDF using WeasyPrint and HTML/CSS templates"""
        try:
            # Get template
            template_name = template_options.get("template_name", "default.html") if template_options else "default.html"
            template = template_env.get_template(template_name)
            
            # Render HTML
            html_content = template.render(cv=cv_data, options=template_options)
            
            # Create a temporary file for the HTML
            with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as tmp_html:
                tmp_html.write(html_content.encode('utf-8'))
                tmp_html_path = tmp_html.name
            
            # Use WeasyPrint to convert HTML to PDF
            html = HTML(filename=tmp_html_path)
            
            # Add any custom styles from template options
            css_styles = []
            if template_options and "custom_css" in template_options:
                css_styles.append(CSS(string=template_options["custom_css"]))
            
            # Generate PDF
            html.write_pdf(output_path, stylesheets=css_styles)
            
            # Clean up temporary file
            os.unlink(tmp_html_path)
            
            logger.info(f"PDF generated at {output_path} using WeasyPrint")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating PDF with WeasyPrint: {str(e)}")
            raise
    
    @staticmethod
    async def generate_pdf_reportlab(cv_data: Dict[str, Any], output_path: str, template_options: Optional[Dict[str, Any]] = None) -> str:
        """Generate a PDF using ReportLab (programmatic approach)"""
        try:
            # Create PDF document
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Define styles
            styles = getSampleStyleSheet()
            title_style = styles['Title']
            heading_style = styles['Heading2']
            normal_style = styles['Normal']
            
            # Customize styles based on template options
            if template_options and "color_scheme" in template_options:
                color_scheme = template_options["color_scheme"]
                if color_scheme == "blue":
                    heading_style.textColor = colors.blue
                elif color_scheme == "green":
                    heading_style.textColor = colors.green
            
            # Build PDF content
            elements = []
            
            # Title
            elements.append(Paragraph(cv_data.get("title", "Curriculum Vitae"), title_style))
            elements.append(Spacer(1, 12))
            
            # Personal Info
            personal_info = cv_data.get("personal_info", {})
            elements.append(Paragraph(f"{personal_info.get('first_name', '')} {personal_info.get('last_name', '')}", styles['Heading1']))
            elements.append(Paragraph(f"Email: {personal_info.get('email', '')}", normal_style))
            elements.append(Paragraph(f"Phone: {personal_info.get('phone', '')}", normal_style))
            elements.append(Paragraph(f"Location: {personal_info.get('location', '')}", normal_style))
            elements.append(Spacer(1, 12))
            
            # Summary
            if "summary" in cv_data and cv_data["summary"]:
                elements.append(Paragraph("Summary", heading_style))
                elements.append(Spacer(1, 6))
                elements.append(Paragraph(cv_data["summary"], normal_style))
                elements.append(Spacer(1, 12))
            
            # Experience
            if "experience" in cv_data and cv_data["experience"]:
                elements.append(Paragraph("Experience", heading_style))
                elements.append(Spacer(1, 6))
                
                for exp in cv_data["experience"]:
                    job_title = exp.get("title", "")
                    company = exp.get("company", "")
                    start_date = exp.get("start_date", "")
                    end_date = "Present" if exp.get("current", False) else exp.get("end_date", "")
                    
                    elements.append(Paragraph(f"<b>{job_title}</b> at {company}", styles['Heading3']))
                    elements.append(Paragraph(f"{start_date} - {end_date}", normal_style))
                    elements.append(Paragraph(exp.get("description", ""), normal_style))
                    elements.append(Spacer(1, 10))
                
                elements.append(Spacer(1, 6))
            
            # Education
            if "education" in cv_data and cv_data["education"]:
                elements.append(Paragraph("Education", heading_style))
                elements.append(Spacer(1, 6))
                
                for edu in cv_data["education"]:
                    degree = edu.get("degree", "")
                    field = edu.get("field_of_study", "")
                    institution = edu.get("institution", "")
                    
                    elements.append(Paragraph(f"<b>{degree}</b> in {field}", styles['Heading3']))
                    elements.append(Paragraph(f"{institution}", normal_style))
                    elements.append(Paragraph(f"{edu.get('start_date', '')} - {edu.get('end_date', '')}", normal_style))
                    if "description" in edu and edu["description"]:
                        elements.append(Paragraph(edu["description"], normal_style))
                    elements.append(Spacer(1, 10))
                
                elements.append(Spacer(1, 6))
            
            # Skills
            if "skills" in cv_data and cv_data["skills"]:
                elements.append(Paragraph("Skills", heading_style))
                elements.append(Spacer(1, 6))
                
                # Create a skills table
                skills_data = []
                row = []
                for i, skill in enumerate(cv_data["skills"]):
                    row.append(skill)
                    if len(row) == 3 or i == len(cv_data["skills"]) - 1:
                        # Fill the row with empty cells if needed
                        while len(row) < 3:
                            row.append("")
                        skills_data.append(row)
                        row = []
                
                if skills_data:
                    skills_table = Table(skills_data, colWidths=[doc.width/3.0]*3)
                    skills_table.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('PADDING', (0, 0), (-1, -1), 6),
                    ]))
                    elements.append(skills_table)
                
                elements.append(Spacer(1, 12))
            
            # Build the PDF
            doc.build(elements)
            
            logger.info(f"PDF generated at {output_path} using ReportLab")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating PDF with ReportLab: {str(e)}")
            raise
    
    @staticmethod
    async def generate_docx(cv_data: Dict[str, Any], output_path: str, template_options: Optional[Dict[str, Any]] = None) -> str:
        """Generate a DOCX document using python-docx"""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.enum.section import WD_ORIENT, WD_SECTION
            from docx.enum.text import WD_LINE_SPACING
            from docx.shared import RGBColor

            # Create a new document
            doc = Document()

            # Set page size and margins (A4, 1" margins)
            section = doc.sections[0]
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

            # Set default font to Arial
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)
            font.color.rgb = RGBColor(0, 0, 0)
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

            # Helper: Add spacing after paragraphs
            def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
                paragraph.paragraph_format.space_before = Pt(before)
                paragraph.paragraph_format.space_after = Pt(after)
                paragraph.paragraph_format.line_spacing = line

            # Helper: Add section heading
            def add_section_heading(text):
                heading = doc.add_paragraph()
                run = heading.add_run(text.upper())
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(0, 0, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_spacing(heading, before=18, after=12, line=1.15)
                # Divider line
                p = doc.add_paragraph()
                p_format = p.paragraph_format
                p_format.space_after = Pt(0)
                p_format.space_before = Pt(0)
                p_format.line_spacing = 1.0
                p._element.get_or_add_pPr().append(OxmlElement('w:pBdr'))
                bdr = p._element.pPr.pBdr
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), 'auto')
                bdr.append(bottom)

            # Helper: Add bullet point
            def add_bullet(text):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(text)
                run.font.size = Pt(11)
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(0, 0, 0)
                set_paragraph_spacing(p, before=0, after=3, line=1.15)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)

            # Helper: Format date
            def format_date(date_str):
                if not date_str:
                    return ''
                try:
                    dt = datetime.strptime(date_str, '%Y-%m-%d')
                    return dt.strftime('%b %Y')
                except Exception:
                    return date_str

            # Candidate Name
            name = f"{cv_data.get('personal_info', {}).get('first_name', '')} {cv_data.get('personal_info', {}).get('last_name', '')}".strip()
            name_p = doc.add_paragraph()
            run = name_p.add_run(name)
            run.bold = True
            run.font.size = Pt(20)
            run.font.name = 'Arial'
            name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(name_p, before=0, after=6, line=1.0)

            # Contact Information
            personal_info = cv_data.get('personal_info', {})
            contact_items = []
            if personal_info.get('phone'):
                contact_items.append(personal_info['phone'])
            if personal_info.get('email'):
                contact_items.append(personal_info['email'].lower())
            if personal_info.get('location'):
                contact_items.append(personal_info['location'])
            contact_str = ' | '.join(contact_items)
            contact_p = doc.add_paragraph(contact_str)
            contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(contact_p, before=0, after=6, line=1.0)

            # LinkedIn (optional)
            if personal_info.get('linkedin'):
                linkedin_p = doc.add_paragraph(personal_info['linkedin'])
                linkedin_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(linkedin_p, before=0, after=6, line=1.0)

            # Summary
            if cv_data.get('summary'):
                add_section_heading('Summary')
                summary_p = doc.add_paragraph(cv_data['summary'])
                set_paragraph_spacing(summary_p, before=0, after=6, line=1.15)

            # Experience
            if cv_data.get('experience'):
                add_section_heading('Experience')
                for exp in cv_data['experience']:
                    # Job title, company, and dates on the same line
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # Job title (dark blue)
                    run = p.add_run(exp.get('title', '').title())
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.name = 'Arial'
                    run.font.color.rgb = RGBColor(0, 51, 153)  # Dark blue
                    # Dates (right-aligned)
                    start = format_date(exp.get('start_date', ''))
                    end = 'Present' if exp.get('current', False) else format_date(exp.get('end_date', ''))
                    if start or end:
                        tab = p.add_run("\t")
                        tab.font.size = Pt(10)
                        tab.font.name = 'Arial'
                        tab.font.color.rgb = RGBColor(0, 0, 0)
                        date_run = p.add_run(f"{start} – {end}")
                        date_run.font.size = Pt(10)
                        date_run.font.name = 'Arial'
                        date_run.font.color.rgb = RGBColor(0, 0, 0)
                    set_paragraph_spacing(p, before=12, after=0, line=1.15)
                    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_ALIGN_PARAGRAPH.RIGHT)
                    # Company name (below, black)
                    company = exp.get('company', '')
                    if company:
                        company_p = doc.add_paragraph(company)
                        company_run = company_p.runs[0]
                        company_run.font.size = Pt(11)
                        company_run.font.name = 'Arial'
                        company_run.font.color.rgb = RGBColor(0, 0, 0)
                        set_paragraph_spacing(company_p, before=0, after=0, line=1.0)
                    # Location (optional, below company)
                    if exp.get('location'):
                        loc_p = doc.add_paragraph(exp['location'])
                        loc_run = loc_p.runs[0]
                        loc_run.font.size = Pt(11)
                        loc_run.font.name = 'Arial'
                        loc_run.font.color.rgb = RGBColor(0, 0, 0)
                        set_paragraph_spacing(loc_p, before=0, after=0, line=1.0)
                    # Bullets
                    desc = exp.get('description', [])
                    if isinstance(desc, str):
                        desc = [line.strip() for line in desc.split('\n') if line.strip()]
                    for i, bullet in enumerate(desc[:5]):
                        add_bullet(bullet)
                    # Space after each role
                    doc.add_paragraph()

            # Education
            if cv_data.get('education'):
                add_section_heading('Education')
                for edu in cv_data['education']:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(f"{edu.get('degree', '').title()} in {edu.get('field_of_study', '')}")
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.name = 'Arial'
                    # Institution
                    inst = edu.get('institution', '')
                    if inst:
                        run2 = p.add_run(f" at {inst}")
                        run2.font.size = Pt(11)
                        run2.font.name = 'Arial'
                    # Dates
                    start = format_date(edu.get('start_date', ''))
                    end = format_date(edu.get('end_date', ''))
                    if start or end:
                        tab = p.add_run("\t")
                        tab.font.size = Pt(10)
                        tab.font.name = 'Arial'
                        date_run = p.add_run(f"{start} – {end}")
                        date_run.font.size = Pt(10)
                        date_run.font.name = 'Arial'
                    set_paragraph_spacing(p, before=12, after=6, line=1.15)
                    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_ALIGN_PARAGRAPH.RIGHT)
                    # Description
                    if edu.get('description'):
                        desc_p = doc.add_paragraph(edu['description'])
                        set_paragraph_spacing(desc_p, before=0, after=6, line=1.15)
                    doc.add_paragraph()

            # Skills
            if cv_data.get('skills'):
                add_section_heading('Skills')
                for skill in cv_data['skills']:
                    add_bullet(skill)
                doc.add_paragraph()

            # Certifications
            if cv_data.get('certifications'):
                add_section_heading('Certifications')
                for cert in cv_data['certifications']:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(cert.get('name', ''))
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.name = 'Arial'
                    issuer = cert.get('issuer', '')
                    date = format_date(cert.get('date', ''))
                    if issuer:
                        run2 = p.add_run(f" by {issuer}")
                        run2.font.size = Pt(11)
                        run2.font.name = 'Arial'
                    if date:
                        tab = p.add_run("\t")
                        tab.font.size = Pt(10)
                        tab.font.name = 'Arial'
                        date_run = p.add_run(f"{date}")
                        date_run.font.size = Pt(10)
                        date_run.font.name = 'Arial'
                    set_paragraph_spacing(p, before=12, after=6, line=1.15)
                    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_ALIGN_PARAGRAPH.RIGHT)
                    if cert.get('description'):
                        desc_p = doc.add_paragraph(cert['description'])
                        set_paragraph_spacing(desc_p, before=0, after=6, line=1.15)
                    doc.add_paragraph()

            # Save the document
            doc.save(output_path)

            logger.info(f"DOCX document generated at {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating DOCX document: {str(e)}")
            raise
    
    @classmethod
    async def generate_document(cls, cv_data: Dict[str, Any], output_path: str, format: str, template_options: Optional[Dict[str, Any]] = None) -> str:
        """Generate a document in the specified format"""
        format = format.lower()
        
        if format == "pdf":
            # Determine which PDF generator to use
            pdf_generator = template_options.get("pdf_generator", "weasyprint") if template_options else "weasyprint"
            
            if pdf_generator == "reportlab":
                return await cls.generate_pdf_reportlab(cv_data, output_path, template_options)
            else:
                return await cls.generate_pdf_weasyprint(cv_data, output_path, template_options)
                
        elif format == "docx":
            return await cls.generate_docx(cv_data, output_path, template_options)
            
        else:
            raise ValueError(f"Unsupported format: {format}") 